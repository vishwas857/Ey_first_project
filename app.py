from flask import Flask, render_template, request, send_file, redirect
import os
from groq import Groq
from dotenv import load_dotenv
import httpx
import io
from docx import Document
from docx.shared import Pt
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
import docx
import markdown
from bs4 import BeautifulSoup
from collections import deque

app = Flask(__name__)

# Load environment variables
load_dotenv()
groq_api_key = os.getenv("GROQ_API_KEY")

http_client = httpx.Client(verify=False)
client = Groq(api_key=groq_api_key, http_client=http_client)

# Conversation buffer (last 5 interactions)
conversation_buffer = deque(maxlen=5)

# Store last essay Markdown
latest_essay = {"markdown": ""}

# Store uploaded PDF/DOCX content + filename
uploaded_document = {"content": None, "filename": None}


# --------------------------
# Extract PDF content
# --------------------------
def extract_pdf_content(pdf_bytes):
    all_text = []
    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")

    for p in range(len(pdf)):
        page = pdf[p]

        text = page.get_text().strip()
        if text:
            all_text.append(text)

        # OCR for images
        for idx, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base = pdf.extract_image(xref)
            img_bytes = base["image"]
            pil_img = Image.open(io.BytesIO(img_bytes))
            ocr_text = pytesseract.image_to_string(pil_img)
            if ocr_text.strip():
                all_text.append(ocr_text)

    pdf.close()
    return "\n\n".join(all_text)


# --------------------------
# Extract DOCX content
# --------------------------
def extract_docx_content(file_bytes):
    file_stream = io.BytesIO(file_bytes)
    doc = docx.Document(file_stream)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


# --------------------------
# Generate essay with memory buffer
# --------------------------
def generate_essay(topic, document_content=None):

    system_message = {
        "role": "system",
        "content": (
            "You are an AI essay writer. Write detailed, structured essays in Markdown format. "
            "Include headings, subheadings, bullet points, and tables where relevant. "
            "If a document is provided, answer ONLY using the document. "
            "If info is missing, say 'Not enough info in the document."
            "if only word is given then generate essay based on that word.'"
        )
    }

    messages = [system_message]
    messages.extend(conversation_buffer)

    if document_content:
        messages.append({"role": "user", "content": f"[DOCUMENT CONTENT]\n{document_content}"})
        messages.append({"role": "user", "content": f"Write an essay on: {topic}"})
    else:
        messages.append({"role": "user", "content": topic})

    response = client.chat.completions.create(
        model="meta-llama/llama-4-scout-17b-16e-instruct",
        messages=messages,
        max_tokens=3500,
        temperature=0
    )

    essay = response.choices[0].message.content.strip()

    conversation_buffer.append({"role": "user", "content": topic})
    conversation_buffer.append({"role": "assistant", "content": essay})

    return essay


# --------------------------
# Convert Markdown → DOCX
# --------------------------
def markdown_to_docx(md_text):
    html = markdown.markdown(md_text, extensions=["tables", "fenced_code", "toc"])
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    for elem in soup.descendants:
        if elem.name == "h1":
            doc.add_heading(elem.get_text(), level=1)
        elif elem.name == "h2":
            doc.add_heading(elem.get_text(), level=2)
        elif elem.name == "h3":
            doc.add_heading(elem.get_text(), level=3)
        elif elem.name == "p":
            doc.add_paragraph(elem.get_text())
        elif elem.name == "li":
            doc.add_paragraph(f"• {elem.get_text()}")
        elif elem.name == "table":
            rows = elem.find_all("tr")
            table = doc.add_table(rows=len(rows),
                                  cols=len(rows[0].find_all(['td', 'th'])))
            for i, row in enumerate(rows):
                cols = row.find_all(['td', 'th'])
                for j, col in enumerate(cols):
                    table.cell(i, j).text = col.get_text()

    return doc


# --------------------------
# ROUTES
# --------------------------
@app.route("/", methods=["GET", "POST"])
def index():
    essay_markdown = ""
    topic = ""

    if request.method == "POST":
        topic = request.form.get("topic", "")
        file = request.files.get("file")

        document_content = None

        if file and file.filename:
            file_bytes = file.read()

            if file.filename.lower().endswith(".pdf"):
                document_content = extract_pdf_content(file_bytes)
            elif file.filename.lower().endswith(".docx"):
                document_content = extract_docx_content(file_bytes)

            uploaded_document["content"] = document_content
            uploaded_document["filename"] = file.filename

        # Generate essay
        if uploaded_document["content"]:
            essay_markdown = generate_essay(topic, uploaded_document["content"])
        else:
            essay_markdown = generate_essay(topic)

        latest_essay["markdown"] = essay_markdown

    return render_template(
        "AIbot.html",
        topic=topic,
        markdown_text=essay_markdown,
        pdf_filename=uploaded_document["filename"]
    )


# --------------------------
# REMOVE FILE ROUTE
# --------------------------
@app.route("/remove_file", methods=["POST"])
def remove_file():
    uploaded_document["content"] = None
    uploaded_document["filename"] = None

    # CLEAR memory buffer so LLM does NOT use document context
    conversation_buffer.clear()

    return redirect("/")


# --------------------------
# DOWNLOAD DOCX
# --------------------------
@app.route("/download")
def download_essay():
    md_text = latest_essay.get("markdown", "")
    if not md_text:
        return "No essay available. Generate one first."

    doc = markdown_to_docx(md_text)
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    return send_file(
        stream,
        as_attachment=True,
        download_name="essay.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


import os

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    app.run(host="0.0.0.0", port=port)

